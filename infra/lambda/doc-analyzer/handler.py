"""
Hromada Partner Document Analyzer — AWS Lambda Handler

Triggered by S3 PutObject events on hromada-partner-docs bucket.
For each uploaded document:
1. Downloads from S3
2. Extracts text (PDF via PyMuPDF, DOCX via python-docx)
3. Classifies document type
4. Extracts cost signals (UAH, USD, EUR amounts)
5. Runs verification checklist against Hromada framework
6. Writes structured analysis JSON to hromada-partner-docs-staging (or a results prefix)
"""
import json
import logging
import os
import re
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import unquote_plus

import boto3
import fitz  # pymupdf
from docx import Document

logger = logging.getLogger()
logger.setLevel(logging.INFO)

s3 = boto3.client("s3")

SOURCE_BUCKET = os.environ.get("SOURCE_BUCKET", "hromada-partner-docs")
RESULTS_BUCKET = os.environ.get("RESULTS_BUCKET", "hromada-partner-docs")
RESULTS_PREFIX = os.environ.get("RESULTS_PREFIX", "_analysis/")


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_pdf(file_path: str) -> dict:
    doc = fitz.open(file_path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text().strip()
        pages.append({"page": i + 1, "text": text})
    metadata = doc.metadata or {}
    full_text = "\n\n".join(p["text"] for p in pages if p["text"])
    result = {
        "page_count": len(doc),
        "metadata": {
            "title": metadata.get("title", ""),
            "author": metadata.get("author", ""),
            "creator": metadata.get("creator", ""),
        },
        "pages": pages,
        "full_text": full_text,
    }
    doc.close()
    return result


def extract_docx(file_path: str) -> dict:
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append({"table_index": i, "rows": rows})
    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "full_text": "\n".join(paragraphs),
    }


# ---------------------------------------------------------------------------
# Classification
# ---------------------------------------------------------------------------

CLASSIFICATION_RULES = [
    ("кошторис", "cost_estimate"),
    ("технічний звіт", "technical_report"),
    ("технічне обстеження", "technical_report"),
    ("рп сес", "design_package"),
    ("рп ", "design_package"),
    ("робочий проект", "design_package"),
    ("_pvs", "project_verification"),
    ("pv", "project_verification"),
    ("акт", "acceptance_act"),
    ("договір", "contract"),
    ("contract", "contract"),
    ("invoice", "invoice"),
    ("рахунок", "invoice"),
]


def classify_document(s3_key: str, text: str) -> str:
    filename = s3_key.lower()
    text_start = text[:1000].lower()

    for pattern, doc_type in CLASSIFICATION_RULES:
        if pattern in filename or pattern in text_start:
            return doc_type
    return "unknown"


# ---------------------------------------------------------------------------
# Cost signal extraction
# ---------------------------------------------------------------------------

CURRENCY_PATTERNS = [
    (r'([\d\s]+[\d.,]+)\s*(?:грн|грив)', "UAH"),
    (r'(?:UAH|грн\.?)\s*([\d\s]+[\d.,]+)', "UAH"),
    (r'([\d\s]+[\d.,]+)\s*(?:EUR|євро|€)', "EUR"),
    (r'([\d\s]+[\d.,]+)\s*(?:USD|дол|\$)', "USD"),
]


def extract_cost_signals(text: str) -> dict:
    amounts = []
    for pattern, currency in CURRENCY_PATTERNS:
        for match in re.finditer(pattern, text):
            raw = match.group(1).replace(" ", "").replace(",", ".")
            try:
                val = float(raw)
                if val > 100:
                    amounts.append({
                        "raw": match.group(0).strip()[:80],
                        "value": val,
                        "currency": currency,
                    })
            except ValueError:
                pass

    # Deduplicate by value+currency
    seen = set()
    unique = []
    for a in amounts:
        key = (a["value"], a["currency"])
        if key not in seen:
            seen.add(key)
            unique.append(a)
    unique.sort(key=lambda x: x["value"], reverse=True)

    return {
        "count": len(unique),
        "amounts": unique[:30],
        "max_amount": unique[0] if unique else None,
        "total_uah": sum(a["value"] for a in unique if a["currency"] == "UAH"),
    }


# ---------------------------------------------------------------------------
# Project info extraction
# ---------------------------------------------------------------------------

def extract_project_info(s3_key: str, text: str) -> dict:
    """Extract structured project info from text content."""
    info = {
        "partner": None,
        "project_name": None,
        "location": None,
        "facility": None,
        "power_kw": None,
    }

    # Partner from S3 key path
    parts = s3_key.split("/")
    if len(parts) > 1:
        info["partner"] = parts[0]
    if len(parts) > 2:
        info["project_name"] = parts[1]

    # Location patterns
    location_match = re.search(
        r'(?:Дніпропетровськ\w+\s+обл[.,]?\s*,?\s*(?:м\.|с\.)\s*\w+|'
        r'Україна,\s*[\w\s]+обл[.,]\s*,?\s*(?:м\.|с\.)\s*[\w\s]+)',
        text[:3000]
    )
    if location_match:
        info["location"] = location_match.group(0).strip()

    # Power (kW)
    power_match = re.search(r'(\d+[\d.,]*)\s*кВт', text[:5000])
    if power_match:
        info["power_kw"] = power_match.group(1)

    return info


# ---------------------------------------------------------------------------
# Verification checklist (based on Hromada Verification Framework)
# ---------------------------------------------------------------------------

VERIFICATION_CHECKLIST = [
    {
        "id": "cost_estimate",
        "label": "Cost estimate (кошторис) provided",
        "description": "Detailed cost breakdown from licensed estimator using AVK-5 or equivalent",
    },
    {
        "id": "design_package",
        "label": "Design/engineering package (РП) provided",
        "description": "Full technical design documentation for the project",
    },
    {
        "id": "technical_report",
        "label": "Technical inspection report provided",
        "description": "Structural/condition assessment by qualified engineer",
    },
    {
        "id": "project_verification",
        "label": "Project verification summary (PV) provided",
        "description": "Partner-level project summary with scope, cost, and co-financing details",
    },
    {
        "id": "avk5_software",
        "label": "Cost estimate generated by AVK-5",
        "description": "Standard Ukrainian construction estimating software — indicates professional preparation",
    },
    {
        "id": "location_identified",
        "label": "Project location clearly identified",
        "description": "Specific address or location referenced in documents",
    },
    {
        "id": "facility_identified",
        "label": "Target facility/object clearly identified",
        "description": "The building or infrastructure being improved is named",
    },
    {
        "id": "usd_cost_available",
        "label": "USD cost estimate available",
        "description": "Cost translated to USD for donor communication",
    },
    {
        "id": "cofinancing_stated",
        "label": "Co-financing commitment stated",
        "description": "Municipality or partner co-financing percentage declared",
    },
    {
        "id": "engineer_identified",
        "label": "Preparing engineer/firm identified",
        "description": "Named professional or firm who prepared the technical documents",
    },
]


def run_verification(s3_key: str, parsed: dict, classification: str, cost_signals: dict, project_info: dict, all_project_docs: list) -> dict:
    """Score the document against the verification checklist.

    all_project_docs is a list of classifications for ALL docs in the same project folder.
    """
    text = parsed.get("full_text", "")
    results = []

    for item in VERIFICATION_CHECKLIST:
        passed = False
        evidence = ""

        if item["id"] == "cost_estimate":
            passed = "cost_estimate" in all_project_docs
            evidence = "Found cost estimate document" if passed else "No cost estimate in project folder"

        elif item["id"] == "design_package":
            passed = "design_package" in all_project_docs
            evidence = "Found design package" if passed else "No design package in project folder"

        elif item["id"] == "technical_report":
            passed = "technical_report" in all_project_docs
            evidence = "Found technical report" if passed else "No technical report in project folder"

        elif item["id"] == "project_verification":
            passed = "project_verification" in all_project_docs
            evidence = "Found PV summary" if passed else "No project verification summary"

        elif item["id"] == "avk5_software":
            passed = "авк" in text[:5000].lower() or "авк-5" in text[:5000].lower()
            evidence = "AVK-5 reference found" if passed else "No AVK-5 reference detected"

        elif item["id"] == "location_identified":
            passed = project_info.get("location") is not None
            evidence = project_info.get("location", "No location found")

        elif item["id"] == "facility_identified":
            passed = any(kw in text[:3000].lower() for kw in ["лікарня", "школа", "вежі", "вежа", "насосн", "котельн", "водонапірн"])
            evidence = "Facility type keyword found" if passed else "No facility keyword detected"

        elif item["id"] == "usd_cost_available":
            max_amt = cost_signals.get("max_amount")
            passed = (max_amt is not None and max_amt.get("currency") == "USD") or any(
                a["currency"] == "USD" for a in cost_signals.get("amounts", [])
            )
            evidence = "USD amount found" if passed else "Only UAH amounts detected"

        elif item["id"] == "cofinancing_stated":
            passed = "співфінанс" in text.lower() or "co-financ" in text.lower() or "cofinanc" in text.lower()
            evidence = "Co-financing reference found" if passed else "No co-financing reference"

        elif item["id"] == "engineer_identified":
            # Look for ФОП (sole proprietor) or company names
            fop_match = re.search(r'ФОП\s+[\w\s.]+', text[:3000])
            if fop_match:
                passed = True
                evidence = fop_match.group(0).strip()
            else:
                evidence = "No engineer/firm identified"

        results.append({
            **item,
            "passed": passed,
            "evidence": evidence,
        })

    passed_count = sum(1 for r in results if r["passed"])
    return {
        "score": f"{passed_count}/{len(results)}",
        "passed": passed_count,
        "total": len(results),
        "percentage": round(passed_count / len(results) * 100, 1),
        "items": results,
    }


# ---------------------------------------------------------------------------
# Lambda handler
# ---------------------------------------------------------------------------

def handler(event, context):
    """Process S3 PutObject events."""
    results = []

    for record in event.get("Records", []):
        bucket = record["s3"]["bucket"]["name"]
        key = unquote_plus(record["s3"]["object"]["key"])

        # Skip analysis output files and non-document files
        if key.startswith(RESULTS_PREFIX):
            logger.info(f"Skipping analysis output: {key}")
            continue

        ext = Path(key).suffix.lower()
        if ext not in (".pdf", ".docx"):
            logger.info(f"Skipping non-document: {key}")
            continue

        logger.info(f"Processing: s3://{bucket}/{key}")

        # Download to temp
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp_path = tmp.name
            s3.download_file(bucket, key, tmp_path)

        try:
            # Extract text
            if ext == ".pdf":
                parsed = extract_pdf(tmp_path)
            else:
                parsed = extract_docx(tmp_path)

            text = parsed.get("full_text", "")

            # Classify
            classification = classify_document(key, text)

            # Extract costs
            cost_signals = extract_cost_signals(text)

            # Extract project info
            project_info = extract_project_info(key, text)

            # List sibling docs in same project folder for cross-reference
            project_prefix = "/".join(key.split("/")[:-1]) + "/"
            sibling_response = s3.list_objects_v2(Bucket=bucket, Prefix=project_prefix)
            sibling_keys = [
                obj["Key"] for obj in sibling_response.get("Contents", [])
                if not obj["Key"].startswith(RESULTS_PREFIX)
            ]

            # For verification, we need classifications of all sibling docs
            # For now, classify by filename only (fast path)
            all_project_doc_types = [classify_document(k, "") for k in sibling_keys]

            # Also check parent folder for PV docs (partner-level)
            partner_prefix = key.split("/")[0] + "/"
            partner_response = s3.list_objects_v2(Bucket=bucket, Prefix=partner_prefix, Delimiter="/")
            partner_files = []
            for obj in partner_response.get("Contents", []):
                if obj["Key"] != partner_prefix:
                    partner_files.append(obj["Key"])
            all_project_doc_types += [classify_document(k, "") for k in partner_files]

            # Run verification
            verification = run_verification(
                key, parsed, classification, cost_signals, project_info, all_project_doc_types
            )

            # Build result (without full_text to keep output manageable)
            analysis = {
                "analyzed_at": datetime.now(timezone.utc).isoformat(),
                "source": f"s3://{bucket}/{key}",
                "s3_key": key,
                "file_type": ext.lstrip("."),
                "classification": classification,
                "project_info": project_info,
                "page_count": parsed.get("page_count"),
                "text_length": len(text),
                "cost_signals": cost_signals,
                "verification": verification,
                "sibling_docs": sibling_keys,
                "text_preview": text[:2000],
            }

            # Write result to S3
            result_key = RESULTS_PREFIX + key.rsplit(".", 1)[0] + ".analysis.json"
            s3.put_object(
                Bucket=RESULTS_BUCKET,
                Key=result_key,
                Body=json.dumps(analysis, ensure_ascii=False, indent=2),
                ContentType="application/json",
            )

            logger.info(
                f"Analysis complete: {key} -> {classification} | "
                f"verification {verification['score']} ({verification['percentage']}%) | "
                f"{cost_signals['count']} cost signals"
            )
            results.append(analysis)

        finally:
            os.unlink(tmp_path)

    # If invoked directly (not from S3 event), return results
    return {
        "statusCode": 200,
        "body": json.dumps({
            "documents_processed": len(results),
            "results": [
                {
                    "key": r["s3_key"],
                    "classification": r["classification"],
                    "verification_score": r["verification"]["score"],
                    "cost_signals": r["cost_signals"]["count"],
                }
                for r in results
            ],
        }, ensure_ascii=False),
    }


# ---------------------------------------------------------------------------
# Local/batch mode: analyze all docs in a bucket
# ---------------------------------------------------------------------------

def analyze_bucket(bucket: str = SOURCE_BUCKET, prefix: str = ""):
    """Run analysis on all documents in a bucket. For local/batch use."""
    paginator = s3.get_paginator("list_objects_v2")
    all_results = []

    for page in paginator.paginate(Bucket=bucket, Prefix=prefix):
        for obj in page.get("Contents", []):
            key = obj["Key"]
            ext = Path(key).suffix.lower()
            if ext not in (".pdf", ".docx") or key.startswith(RESULTS_PREFIX):
                continue

            fake_event = {
                "Records": [{
                    "s3": {
                        "bucket": {"name": bucket},
                        "object": {"key": key},
                    }
                }]
            }
            result = handler(fake_event, None)
            all_results.append(json.loads(result["body"]))

    return all_results


if __name__ == "__main__":
    import sys
    bucket = sys.argv[1] if len(sys.argv) > 1 else SOURCE_BUCKET
    prefix = sys.argv[2] if len(sys.argv) > 2 else ""
    print(f"Analyzing all documents in s3://{bucket}/{prefix}")
    results = analyze_bucket(bucket, prefix)
    print(json.dumps(results, ensure_ascii=False, indent=2))
